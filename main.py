import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import concurrent.futures
import fitz  # PyMuPDF
import webbrowser
from pathlib import Path
import re
import threading
import csv
import datetime
from difflib import SequenceMatcher

# ==========================================
# CONFIGURATION
# ==========================================

SUPPORTED_EXTS = {
    'pdf': ['.pdf'],
    'word': ['.docx', '.doc'],
    'text': ['.txt', '.py', '.c', '.cpp', '.h', '.java', '.md', '.json', '.xml', '.csv']
}

CLEAN_TRANS = str.maketrans('', '', ' \n\t\r')

try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ==========================================
# CORE TEXT EXTRACTION (For Keyword Search)
# ==========================================

def extract_text_from_file(path):
    ext = os.path.splitext(path)[1].lower()
    text = ""
    try:
        if ext in SUPPORTED_EXTS['pdf']:
            with fitz.open(path) as doc:
                for page in doc: text += page.get_text() + "\n"
        elif ext in SUPPORTED_EXTS['word'] and HAS_DOCX:
            doc = docx.Document(path)
            for p in doc.paragraphs: text += p.text + "\n"
        else:
            try:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except:
                with open(path, 'rb') as f:
                    text = f.read().decode('utf-8', errors='ignore')
    except:
        return ""
    return text


# ==========================================
# WORKERS
# ==========================================

def _check_match(raw_text, query, query_rev, use_regex):
    if not raw_text: return (False, "")
    text_lower = raw_text.lower()
    found = False
    match_term = ""

    if use_regex:
        try:
            match = re.search(query, text_lower)
            if match: found = True; match_term = match.group(0)
        except:
            pass
    else:
        text_compressed = text_lower.translate(CLEAN_TRANS)
        if query in text_compressed:
            found = True; match_term = query
        elif query_rev and (query_rev in text_compressed):
            found = True; match_term = query_rev

    if found:
        term = match_term if match_term else query
        start = text_lower.find(term)
        if start == -1 and not use_regex: start = 0
        if start != -1:
            s = max(0, start - 40);
            e = min(len(text_lower), start + 40)
            snip = raw_text[s:e].replace("\n", " ").strip()
            return (True, f"...{snip}...")
        return (True, "Match found")
    return (False, "")


def worker_search_file(task_data):
    path, query, query_rev, use_regex = task_data
    full_text = extract_text_from_file(path)
    res = _check_match(full_text, query, query_rev, use_regex)
    if res[0]: return (True, path, "Text", res[1])
    return (False, path, 0, "")


def get_tokens(text):
    words = re.findall(r'\w+', text.lower())
    return set(w for w in words if len(w) > 3)


def worker_similarity_scan(task_data):
    target_path, ref_tokens = task_data
    target_text = extract_text_from_file(target_path)
    if not target_text: return (False, target_path, 0)
    target_tokens = get_tokens(target_text)
    if not target_tokens: return (False, target_path, 0)

    intersection = len(ref_tokens.intersection(target_tokens))
    union = len(ref_tokens.union(target_tokens))
    score = (intersection / union) * 100 if union > 0 else 0

    if score > 5.0: return (True, target_path, round(score, 1))
    return (False, target_path, 0)


# ==========================================
# DEEP INSPECTION LOGIC (Page Aware)
# ==========================================

# ==========================================
# DEEP INSPECTION LOGIC (Page Aware)
# ==========================================

class DeepInspector:
    @staticmethod
    def extract_pages(path):
        """
        Returns a list of tuples: [(page_num, text), (page_num, text)...]
        (This method remains unchanged as it handles file reading)
        """
        ext = os.path.splitext(path)[1].lower()
        pages = []

        try:
            if ext in SUPPORTED_EXTS['pdf']:
                with fitz.open(path) as doc:
                    for i, page in enumerate(doc):
                        pages.append((str(i + 1), page.get_text()))

            elif ext in SUPPORTED_EXTS['word'] and HAS_DOCX:
                # Word docs don't have real pages, use Paragraph clusters or just "1"
                doc = docx.Document(path)
                full_text = "\n".join([p.text for p in doc.paragraphs])
                pages.append(("1", full_text))

            else:
                # Text files
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    pages.append(("1", f.read()))

        except:
            pass
        return pages

    @staticmethod
    def parse_text_chunks_with_location(path):
        """
        Splits the document into chunks based on two or more newlines.
        Maps each chunk to its starting page number.
        Returns: [ {'id': '1', 'text': '...', 'page': '5'}, ... ]
        """
        pages_data = DeepInspector.extract_pages(path)
        parsed_items = []
        chunk_id = 1

        for page_num, page_text in pages_data:
            # Use regex to split text by two or more consecutive newline characters
            # This creates chunks based on logical paragraph breaks
            chunks = re.split(r'\n{2,}', page_text.strip())

            for chunk in chunks:
                cleaned_chunk = chunk.strip()
                if len(cleaned_chunk) > 10:  # Minimum length filter for meaningful chunks
                    parsed_items.append({
                        'id': str(chunk_id),
                        'text': cleaned_chunk,
                        # For simplicity, assign the entire chunk to the page it starts on.
                        'page': page_num
                    })
                    chunk_id += 1

        return parsed_items

    @staticmethod
    def compare_structure(ref_path, target_path):
        """
        Compares text chunks between the reference and target documents.
        The function name is kept for backward compatibility with the GUI call.
        """
        # --- Using the new chunk parsing function ---
        ref_data = DeepInspector.parse_text_chunks_with_location(ref_path)
        tgt_data = DeepInspector.parse_text_chunks_with_location(target_path)
        # ------------------------------------------

        results = []

        for ref_item in ref_data:
            # Filter short chunks
            if len(ref_item['text']) < 50: continue

            best_score = 0.0
            best_match_page = "-"

            for tgt_item in tgt_data:
                if len(tgt_item['text']) < 50: continue

                # Calculate the similarity ratio
                # You might want to strip common whitespace and case-fold for better matches
                ref_text_clean = ref_item['text'].lower().translate(CLEAN_TRANS)
                tgt_text_clean = tgt_item['text'].lower().translate(CLEAN_TRANS)

                ratio = SequenceMatcher(None, ref_text_clean, tgt_text_clean).ratio() * 100

                if ratio > best_score:
                    best_score = ratio
                    best_match_page = tgt_item['page']

            if best_score > 15:  # Lower the threshold slightly for general text chunk comparison
                results.append({
                    'ref_page': ref_item['page'],
                    'tgt_page': best_match_page,
                    'score': round(best_score, 1),
                    'preview': ref_item['text'][:100].replace('\n', ' ') + "..."
                })

        # Sort by score
        results.sort(key=lambda x: x['score'], reverse=True)
        return results


# ==========================================
# LOGIC
# ==========================================

class SearchLogic:
    def __init__(self):
        self.stop_flag = False

    def get_files(self, folder):
        valid = set(e for sub in SUPPORTED_EXTS.values() for e in sub)
        files = []
        myself = os.path.abspath(__file__)
        for r, _, fs in os.walk(folder):
            for f in fs:
                full = os.path.join(r, f)
                if os.path.abspath(full) == myself: continue
                if os.path.splitext(f)[1].lower() in valid:
                    files.append(full)
        return files

    def run_keyword_search(self, files, query, regex, cbs):
        self.stop_flag = False
        q_c = query if regex else query.lower().replace(" ", "").replace("\t", "")
        q_r = None if regex else q_c[::-1]
        tasks = [(f, q_c, q_r, regex) for f in files]
        self._run_pool(worker_search_file, tasks, cbs)

    def run_similarity_search(self, files, ref_file, cbs):
        self.stop_flag = False
        ref_text = extract_text_from_file(ref_file)
        ref_tokens = get_tokens(ref_text)
        if not ref_tokens: return
        ref_abs = os.path.abspath(ref_file)
        tasks = [(f, ref_tokens) for f in files if os.path.abspath(f) != ref_abs]
        self._run_pool(worker_similarity_scan, tasks, cbs)

    def _run_pool(self, worker, tasks, cbs):
        total = len(tasks)
        count = 0
        matches = 0
        with concurrent.futures.ProcessPoolExecutor() as ex:
            fs = {ex.submit(worker, t): t[0] for t in tasks}
            try:
                for f in concurrent.futures.as_completed(fs):
                    if self.stop_flag: ex.shutdown(wait=False); break
                    count += 1
                    try:
                        res = f.result()
                        if res[0]:
                            matches += 1
                            if cbs.get('on_match'): cbs['on_match'](res)
                    except:
                        pass
                    if cbs.get('on_prog') and (count % 5 == 0 or count == total):
                        cbs['on_prog'](int(count / total * 100), os.path.basename(fs[f]))
            except KeyboardInterrupt:
                ex.shutdown(wait=False)
        if cbs.get('on_done'): cbs['on_done'](matches)

    def stop(self):
        self.stop_flag = True


# ==========================================
# GUI
# ==========================================

class SearchGUI:
    def __init__(self, root):
        self.root = root
        self.logic = SearchLogic()
        self.root.title("Deep Content Searcher")
        self.root.geometry("1250x800")
        self.last_search_type = None
        self.last_search_query = None
        self._init_ui()
        self.selected_folder = ""
        self.is_searching = False

    def _init_ui(self):
        top = tk.Frame(self.root, pady=10);
        top.pack(fill="x", padx=10)
        self.lbl_path = tk.Label(top, text="No folder selected", fg="gray", anchor="w")
        self.lbl_path.pack(side="left", fill="x", expand=True)
        self.btn_export = tk.Button(top, text="Export CSV", command=self.export_csv, state="disabled")
        self.btn_export.pack(side="right", padx=5)
        tk.Button(top, text="Select Archive", command=self.browse_folder).pack(side="right")

        frm = tk.Frame(self.root, pady=5);
        frm.pack(fill="x", padx=10)

        lf = tk.LabelFrame(frm, text="Keyword Search");
        lf.pack(side="left", fill="both", expand=True, padx=5)
        self.entry = tk.Entry(lf);
        self.entry.pack(side="left", fill="x", expand=True, padx=5, pady=5)
        self.entry.bind('<Key>', self.handle_hotkeys);
        self.entry.bind('<Return>', lambda e: self.start_keyword())
        self.var_reg = tk.BooleanVar()
        tk.Checkbutton(lf, text="Regex", variable=self.var_reg).pack(side="left")
        self.btn_key = tk.Button(lf, text="Search", command=self.start_keyword, bg="#d1e7dd")
        self.btn_key.pack(side="left", padx=5)

        rf = tk.LabelFrame(frm, text="Overlap Finder");
        rf.pack(side="right", fill="both", expand=True, padx=5)
        tk.Label(rf, text="Select file to find reuse:").pack(side="left", padx=5)
        self.btn_sim = tk.Button(rf, text="Find Similar...", command=self.start_similarity, bg="#ffe6cc")
        self.btn_sim.pack(side="right", padx=5)

        pf = tk.Frame(self.root);
        pf.pack(fill="x", padx=10, pady=5)
        self.lbl_stat = tk.Label(pf, text="Ready", anchor="w", fg="blue");
        self.lbl_stat.pack(fill="x")
        self.prog = ttk.Progressbar(pf, orient="horizontal", mode="determinate");
        self.prog.pack(fill="x")

        paned = tk.PanedWindow(self.root, orient=tk.VERTICAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        frame_table = tk.Frame(paned)
        paned.add(frame_table, height=400)

        self.tree = ttk.Treeview(frame_table, columns=("name", "dir", "loc", "ctx", "path"), show="headings")
        self.tree.tag_configure('high', background='#d4edda');
        self.tree.tag_configure('mid', background='#fff3cd')
        self.tree.tag_configure('odd', background='#f8f9fa');
        self.tree.tag_configure('even', background='#ffffff')

        self.tree.heading("name", text="File Name", command=lambda: self.sort("name", False))
        self.tree.heading("dir", text="Directory", command=lambda: self.sort("dir", False))
        self.tree.heading("loc", text="Loc/Score", command=lambda: self.sort("loc", False))
        self.tree.heading("ctx", text="Context", command=lambda: self.sort("ctx", False))

        self.tree.column("name", width=200);
        self.tree.column("dir", width=200);
        self.tree.column("loc", width=80, anchor="center")
        self.tree.column("ctx", width=500);
        self.tree.column("path", width=0, stretch=False)

        sc = tk.Scrollbar(frame_table, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=sc.set);
        self.tree.pack(side="left", fill="both", expand=True);
        sc.pack(side="right", fill="y")

        self.tree.bind('<Double-1>', self.on_open);
        self.tree.bind('<Button-3>', self.on_menu)
        self.tree.bind('<<TreeviewSelect>>', self.on_select_row)

        frame_prev = tk.LabelFrame(paned, text="Instant Preview");
        paned.add(frame_prev, minsize=150)
        self.txt_prev = tk.Text(frame_prev, wrap="word", font=("Consolas", 10), state="disabled", bg="#fcfcfc")
        sc_p = tk.Scrollbar(frame_prev, orient="vertical", command=self.txt_prev.yview)
        self.txt_prev.configure(yscrollcommand=sc_p.set);
        self.txt_prev.pack(side="left", fill="both", expand=True);
        sc_p.pack(side="right", fill="y")

        self.menu = tk.Menu(self.root, tearoff=0)
        self.menu.add_command(label="Copy Full Path", command=self.copy_path)

    def handle_hotkeys(self, e):
        ctrl = (e.state & 4) or (e.state & 131072)
        if ctrl and e.keycode in [65, 67, 86, 88]:
            action = {65: lambda: (e.widget.select_range(0, tk.END), e.widget.icursor(tk.END)),
                      67: lambda: e.widget.event_generate("<<Copy>>"),
                      86: lambda: e.widget.event_generate("<<Paste>>"),
                      88: lambda: e.widget.event_generate("<<Cut>>")}
            action[e.keycode]();
            return "break"

    def browse_folder(self):
        f = filedialog.askdirectory()
        if f: self.selected_folder = f; self.lbl_path.config(text=f, fg="black")

    def start_keyword(self):
        if not self.check_ready(): return
        q = self.entry.get();
        if not q: return messagebox.showwarning("Warn", "Enter text.")
        self.last_search_type = "Keyword";
        self.last_search_query = q
        self.prep_search("Indexing...");
        self.tree.heading("loc", text="Location")
        threading.Thread(target=self.thread_key, args=(q,), daemon=True).start()

    def start_similarity(self):
        if not self.check_ready(): return
        ref = filedialog.askopenfilename(title="Select Reference File")
        if not ref: return
        self.last_search_type = "Similarity";
        self.last_search_query = ref
        self.prep_search(f"Scanning against '{os.path.basename(ref)}'...")
        self.tree.heading("loc", text="Score (%)")
        threading.Thread(target=self.thread_sim, args=(ref,), daemon=True).start()

    def prep_search(self, msg):
        if self.is_searching: self.logic.stop(); return
        self.is_searching = True;
        self.btn_export.config(state="disabled");
        self.btn_key.config(state="disabled");
        self.btn_sim.config(state="disabled")
        for i in self.tree.get_children(): self.tree.delete(i)
        self.lbl_stat.config(text=msg);
        self.root.update()

    def check_ready(self):
        if not self.selected_folder: messagebox.showwarning("Warn", "Select Archive Folder first."); return False
        return True

    def thread_key(self, q):
        files = self.logic.get_files(self.selected_folder)
        self.run_common(files, lambda f, cbs: self.logic.run_keyword_search(f, q, self.var_reg.get(), cbs))

    def thread_sim(self, ref):
        files = self.logic.get_files(self.selected_folder)
        self.run_common(files, lambda f, cbs: self.logic.run_similarity_search(f, ref, cbs))

    def run_common(self, files, func):
        self.root.after(0, lambda: self.prog.configure(maximum=len(files)))
        cbs = {'on_match': lambda r: self.root.after(0, lambda: self.add_row(r)),
               'on_prog': lambda p, n: self.root.after(0, lambda: self.update_ui(p, n)),
               'on_done': lambda c: self.root.after(0, lambda: self.done(c))}
        func(files, cbs)

    def add_row(self, r):
        full = r[1];
        name = os.path.basename(full)
        try:
            directory = os.path.relpath(os.path.dirname(full), self.selected_folder)
        except:
            directory = os.path.dirname(full)

        row_tag = 'even' if len(self.tree.get_children()) % 2 == 0 else 'odd'
        if len(r) != 4:  # Similarity
            if r[2] >= 80:
                row_tag = 'high'
            elif r[2] >= 40:
                row_tag = 'mid'
            vals = (name, directory, f"{r[2]}%", "Content Overlap", full)
        else:
            vals = (name, directory, r[2], r[3], full)

        self.tree.insert("", "end", values=vals, tags=(row_tag,))

    def update_ui(self, p, n):
        self.prog['value'] = int(p * self.prog['maximum'] / 100);
        self.lbl_stat.config(text=f"Scanning: {n[:40]}...")

    def done(self, c):
        self.is_searching = False;
        self.btn_key.config(state="normal");
        self.btn_sim.config(state="normal")
        if c > 0: self.btn_export.config(state="normal")
        self.lbl_stat.config(text=f"Search Completed. Found {c} matches.");
        self.prog['value'] = 0

        if self.last_search_type == "Similarity":
            # The 'loc' column holds the score (%)
            # Sort by 'loc' (score) in descending order (reverse=True)
            self.sort("loc", True)

    def export_csv(self):
        if not self.tree.get_children(): return
        path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if not path: return
        try:
            with open(path, 'w', newline='', encoding='utf-8-sig') as f:
                w = csv.writer(f)
                w.writerow(["Report:", datetime.datetime.now(), "Root:", self.selected_folder])
                w.writerow(["Type:", self.last_search_type, "Query:", self.last_search_query, ""]);
                w.writerow(["File", "Dir", "Loc/Score", "Context", "Path"])
                for i in self.tree.get_children(): w.writerow(self.tree.item(i)['values'])
            messagebox.showinfo("OK", "Exported.")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def on_select_row(self, e):
        sel = self.tree.selection()
        if not sel: return
        vals = self.tree.item(sel[0], "values")
        loc = str(vals[2]);
        path = vals[4]
        arg = None if "%" in loc else loc
        threading.Thread(target=self.fetch_prev, args=(path, arg), daemon=True).start()

    def fetch_prev(self, path, loc):
        text = extract_text_from_file(path)
        self.root.after(0, lambda: self.update_prev(text))

    def update_prev(self, text):
        self.txt_prev.config(state="normal");
        self.txt_prev.delete("1.0", tk.END);
        self.txt_prev.insert("1.0", text if text else "No text.");
        self.txt_prev.config(state="disabled")

    def on_open(self, e):
        if not self.tree.identify_row(e.y): return
        vals = self.tree.item(self.tree.selection()[0], "values")
        try:
            if vals[2].isdigit() and vals[4].endswith('.pdf'):
                webbrowser.open(f"{Path(vals[4]).as_uri()}#page={vals[2]}")
            else:
                os.startfile(vals[4])
        except:
            pass

    def on_menu(self, e):
        row = self.tree.identify_row(e.y)
        if row:
            self.tree.selection_set(row)
            self.menu.delete(1, tk.END)
            if self.last_search_type == "Similarity":
                self.menu.add_command(label="Deep Inspect Chunks...", command=self.launch_deep_inspection)
            self.menu.tk_popup(e.x_root, e.y_root)

    def copy_path(self):
        sel = self.tree.selection()
        if sel: self.root.clipboard_clear(); self.root.clipboard_append(
            self.tree.item(sel[0], "values")[4]); self.root.update()

    def sort(self, col, rev):
        l = [(self.tree.set(k, col), k) for k in self.tree.get_children('')]
        try:
            l.sort(key=lambda t: float(t[0].replace('%', '')), reverse=rev)
        except:
            l.sort(key=lambda t: t[0].lower(), reverse=rev)
        for i, (v, k) in enumerate(l): self.tree.move(k, '', i)
        self.tree.heading(col, command=lambda: self.sort(col, not rev))

    def launch_deep_inspection(self):
        sel = self.tree.selection()
        if not sel: return
        target_file = self.tree.item(sel[0], "values")[4]
        ref_file = self.last_search_query
        InspectWindow(self.root, ref_file, target_file)


class InspectWindow:
    def __init__(self, parent, ref_path, target_path):
        self.win = tk.Toplevel(parent);
        self.win.title("Deep Chunk Comparison");
        self.win.geometry("1000x600")
        self.ref_path = ref_path;
        self.target_path = target_path
        tk.Label(self.win, text=f"Ref: {os.path.basename(ref_path)}", fg="blue").pack(pady=2)
        tk.Label(self.win, text=f"Target: {os.path.basename(target_path)}", fg="red").pack(pady=2)
        self.lbl_stat = tk.Label(self.win, text="Analyzing...");
        self.lbl_stat.pack()

        # UPDATED COLUMNS: Page (Source) | Page (Target) | Score | Preview
        self.tree = ttk.Treeview(self.win, columns=("p1", "p2", "scr", "txt"), show="headings")
        self.tree.heading("p1", text="Page (Source)");
        self.tree.heading("p2", text="Page (Target)")
        self.tree.heading("scr", text="Score %");
        self.tree.heading("txt", text="Matched Text Preview")

        self.tree.column("p1", width=100, anchor="center");
        self.tree.column("p2", width=100, anchor="center")
        self.tree.column("scr", width=80, anchor="center");
        self.tree.column("txt", width=600)

        sc = tk.Scrollbar(self.win, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=sc.set);
        self.tree.pack(side="left", fill="both", expand=True);
        sc.pack(side="right", fill="y")

        self.tree.tag_configure('high', background='#d4edda');
        self.tree.tag_configure('mid', background='#fff3cd')

        # --- ENHANCEMENT: Bind double-click for page navigation ---
        self.tree.bind('<Double-1>', self.on_open)
        # ---------------------------------------------------------

        threading.Thread(target=self.run_analysis, daemon=True).start()

    def run_analysis(self):
        results = DeepInspector.compare_structure(self.ref_path, self.target_path)
        self.win.after(0, lambda: self.show_results(results))

    def show_results(self, results):
        self.lbl_stat.config(text=f"Analysis complete. Found {len(results)} comparisons.")
        for r in results:
            tag = 'high' if r['score'] > 80 else ('mid' if r['score'] > 50 else '')
            # Store full path in the item's values (not displayed but accessible)
            self.tree.insert("", "end",
                             values=(r['ref_page'], r['tgt_page'], f"{r['score']}%", r['preview'], self.ref_path,
                                     self.target_path),
                             tags=(tag,))

    def on_open(self, e):
        """Opens the source or target file at the matched page number on double-click."""
        item_id = self.tree.identify_row(e.y)
        if not item_id: return

        # 1. Get column that was clicked
        column_clicked = self.tree.identify_column(e.x)

        # 2. Get item values: (p1, p2, scr, txt, ref_path, target_path)
        vals = self.tree.item(item_id, "values")

        file_to_open = None
        page_num = None

        # Determine which file/page to open based on the clicked column
        # #0 is invisible, #1 is p1 (Source Page), #2 is p2 (Target Page)
        if column_clicked == '#1':  # Page (Source) column clicked
            file_to_open = vals[4]  # ref_path
            page_num = vals[0]  # ref_page
        elif column_clicked == '#2':  # Page (Target) column clicked
            file_to_open = vals[5]  # target_path
            page_num = vals[1]  # tgt_page
        else:
            return  # Ignore clicks on other columns

        if file_to_open and page_num and page_num.isdigit():
            try:
                # Open PDF directly to page if it's a PDF
                if file_to_open.lower().endswith('.pdf'):
                    # Use platform-independent URI scheme with fragment for page number
                    webbrowser.open(f"{Path(file_to_open).as_uri()}#page={page_num}")
                else:
                    # Open non-PDF file normally
                    os.startfile(file_to_open)
            except Exception as ex:
                messagebox.showerror("Open Error", f"Could not open file:\n{ex}")

if __name__ == "__main__":
    import multiprocessing

    multiprocessing.freeze_support()
    root = tk.Tk()
    app = SearchGUI(root)
    root.mainloop()